# -*- coding: utf-8 -*-
"""
Build the Information Fusion submission as a Word (.docx) document.

- Body font: Arial. Figures embedded at 190/140/85 mm (Elsevier double/1.5/single col).
- Display equations are rendered to 600-DPI images with matplotlib mathtext, so every
  symbol, sub/superscript, and fraction renders exactly (no font-substitution garble).
- Figures are read from ./figs (produced by figures.py); equation images go to ./eqs.

mathtext gotchas (do not reintroduce): no \\lVert (use \\|), no \\lvert (use |),
no \\big (use \\left/\\right), no \\bigoplus (use \\oplus), no \\Pr (use \\mathrm{P}),
no \\varnothing (use \\emptyset), no \\mathbb. A literal " inside r"..." ends the string.

Run:  python build_docx.py
Out:  ./CMPSA_InformationFusion.docx
"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
from PIL import Image

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
FIGS = os.path.join(HERE, "figs")
EQS = os.path.join(HERE, "eqs")
os.makedirs(EQS, exist_ok=True)

ARIAL = "Arial"
BODY_PT = 10
GRAY = RGBColor(0x55, 0x55, 0x55)


def _register_arial():
    for fn in ("arial.ttf", "arialbd.ttf", "ariali.ttf"):
        p = os.path.join(r"C:\Windows\Fonts", fn)
        if os.path.exists(p):
            try:
                font_manager.fontManager.addfont(p)
            except Exception:
                pass


_register_arial()
plt.rcParams.update({
    "font.family": "Arial", "font.sans-serif": ["Arial", "DejaVu Sans"],
    "mathtext.fontset": "stixsans", "text.color": "#111111",
})


# ---------------------------------------------------------------- equations
def render_eq(tag, tex, fontsize=17):
    fig = plt.figure(figsize=(0.1, 0.1))
    fig.text(0.5, 0.5, tex, ha="center", va="center", fontsize=fontsize)
    out = os.path.join(EQS, f"eq_{tag}.png")
    fig.savefig(out, dpi=600, transparent=True, bbox_inches="tight", pad_inches=0.03)
    plt.close(fig)
    return out


def eq_width_mm(path, cap=132.0):
    w_px, _ = Image.open(path).size
    return min(w_px / 600.0 * 25.4, cap)


# ---------------------------------------------------------------- docx helpers
def set_cell_font(cell, bold=False, size=8):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        for r in p.runs:
            r.font.name = ARIAL
            r.font.size = Pt(size)
            r.font.bold = bold
            r._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)


def add_body(doc, text, size=BODY_PT, space_after=6, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = ARIAL
    r.font.size = Pt(size)
    r.italic = italic
    r._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)
    return p


def add_runs(doc, segments, size=BODY_PT, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for text, style in segments:
        r = p.add_run(text)
        r.font.name = ARIAL
        r.font.size = Pt(size)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)
        if style == "i":
            r.italic = True
        elif style == "b":
            r.bold = True
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = ARIAL
    r.bold = True
    r.font.size = Pt(12 if level == 1 else 10.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)
    return p


def add_equation(doc, tag, tex, fontsize=17):
    path = render_eq(tag, tex, fontsize)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Mm(eq_width_mm(path)))
    num = p.add_run(f"    ({tag})")
    num.font.name = ARIAL
    num.font.size = Pt(BODY_PT)
    return p


def add_figure(doc, fname, width_mm, caption_segments):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(FIGS, fname), width=Mm(width_mm))
    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(8)
    for text, style in caption_segments:
        r = cap.add_run(text)
        r.font.name = ARIAL
        r.font.size = Pt(8.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)
        if style == "b":
            r.bold = True
        elif style == "i":
            r.italic = True


def add_table(doc, caption_segments, headers, rows, colwidths_mm=None, note=None):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(2)
    for text, style in caption_segments:
        r = cap.add_run(text)
        r.font.name = ARIAL
        r.font.size = Pt(8.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)
        if style == "b":
            r.bold = True
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_font(hdr[i], bold=True, size=8)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_cell_font(cells[i], bold=False, size=8)
    if colwidths_mm:
        for i, wmm in enumerate(colwidths_mm):
            for r in t.rows:
                r.cells[i].width = Mm(wmm)
    if note:
        n = doc.add_paragraph()
        n.paragraph_format.space_after = Pt(8)
        rn = n.add_run(note)
        rn.font.name = ARIAL
        rn.font.size = Pt(7.5)
        rn.italic = True
        rn.font.color.rgb = GRAY
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ================================================================= BUILD
def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = ARIAL
    normal.font.size = Pt(BODY_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)
    s = doc.sections[0]
    s.left_margin = Mm(25); s.right_margin = Mm(25)
    s.top_margin = Mm(22); s.bottom_margin = Mm(22)

    # ---- title ----
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("Hierarchical Grounding Fusion for Detecting and Mitigating "
                    "Object, Attribute, and Relation Hallucinations in Multimodal "
                    "Large Language Models")
    tr.font.name = ARIAL; tr.font.size = Pt(15); tr.bold = True
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Prepared for Information Fusion (Elsevier) — anonymous submission")
    sr.font.name = ARIAL; sr.font.size = Pt(9); sr.italic = True; sr.font.color.rgb = GRAY

    # ---- highlights ----
    add_heading(doc, "Highlights", 1)
    for h in HIGHLIGHTS:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        b = p.add_run("•  " + h); b.font.name = ARIAL; b.font.size = Pt(9.5)

    add_heading(doc, "Abstract", 1)
    add_body(doc, ABSTRACT, size=9.5)
    add_runs(doc, [("Keywords: ", "b"), (KEYWORDS, "")], size=9.5)

    # ---- 1 Introduction ----
    add_heading(doc, "1. Introduction", 1)
    for para in INTRO:
        add_body(doc, para)
    add_figure(doc, "fig1_framework.png", 190,
               [("Fig. 1. ", "b"),
                ("Overall research framework, read top to bottom. Problem: the three "
                 "kinds of hallucination and the open question of whether detection "
                 "transfers to mitigation. Method: hierarchical grounding fusion — "
                 "three matched detectors and three uses, training-free and deployable "
                 "as a wrapper. Validation: full-scale detection, mitigation, and "
                 "discrimination results on four backbones, including the negative "
                 "decoder-steering result. Findings and deliverables: the "
                 "detection–mitigation gap and the residual principle, shipped as the "
                 "pip-installable hgfusion plug-in, the HalluProbe-VL diagnostic, and "
                 "open code.", "")])

    # ---- 2 Related work ----
    add_heading(doc, "2. Related work", 1)
    for para in RELATED:
        add_body(doc, para)
    add_table(doc,
              [("Table 1. ", "b"),
               ("Positioning against representative hallucination methods. "
                "O/A/R = object/attribute/relation; ✓ yes, ~ partial, – no. For our "
                "method, attribute and relation are covered at the detection level and "
                "the revision is validated for objects (Section 6).", "")],
              ["Method", "Paradigm", "O", "A", "R", "Detect", "Revise",
               "Train-free", "Multi-backbone", "Gap analysis"],
              TABLE1_ROWS,
              colwidths_mm=[26, 26, 7, 7, 7, 13, 13, 15, 20, 18])

    # ---- 3 Method ----
    add_heading(doc, "3. Method", 1)
    add_body(doc, METHOD["intro"])
    add_figure(doc, "fig2_principle.png", 190,
               [("Fig. 2. ", "b"),
                ("Structure of the algorithm, read left to right. One (image, output) "
                 "pair is processed in a single pass: the output is decomposed into "
                 "claims (Eq. 1); each claim is scored by its matched signal — CLIP "
                 "cosine for objects, region-crop contrast for attributes, box geometry "
                 "for relations (Eqs. 4–7). All three scores feed the threshold test "
                 "(Eq. 3) and the evidence report, while the object score alone drives "
                 "the revise decision (Eqs. 8–9) and the calibrated decision fusion "
                 "(Eqs. 12–13), yielding a verified caption or a calibrated answer.", "")])

    add_heading(doc, "3.1. Problem formulation", 2)
    add_body(doc, METHOD["formulation"][0])
    add_equation(doc, "1", r"$M(y)=M_{\mathrm{obj}}(y)\,\cup\,M_{\mathrm{attr}}(y)\,\cup\,M_{\mathrm{rel}}(y)$")
    add_body(doc, METHOD["formulation"][1])

    add_heading(doc, "3.2. Detection as a hypothesis test", 2)
    add_body(doc, METHOD["hypothesis"][0])
    add_equation(doc, "2",
        r"$\Lambda(c)=\log\frac{p\left(g(c)\mid H_0\right)}{p\left(g(c)\mid H_1\right)},"
        r"\qquad \mathrm{decide}\;H_1\;\;\mathrm{iff}\;\;\Lambda(c)<\eta$")
    add_body(doc, METHOD["hypothesis"][1])
    add_equation(doc, "3",
        r"$\Lambda\;\mathrm{nondecreasing\;in}\;g\quad\Rightarrow\quad"
        r"\hat{h}(c)=\mathbf{1}\!\left[\,g(c)<\tau\,\right]$")
    add_body(doc, METHOD["hypothesis"][2])

    add_heading(doc, "3.3. Hierarchical grounding signals", 2)
    add_body(doc, METHOD["grounding"][0])
    add_equation(doc, "4",
        r"$g_{\mathrm{obj}}(I,o)=\frac{\langle\, \phi_v(I),\;\phi_t(\pi_o)\,\rangle}"
        r"{\|\phi_v(I)\|\;\|\phi_t(\pi_o)\|}$")
    add_body(doc, METHOD["grounding"][1])
    add_equation(doc, "5",
        r"$g_{\mathrm{attr}}(I,o,a)=\cos\!\left(\phi_v(R_o),\phi_t(\pi_{a,o})\right)"
        r"-\cos\!\left(\phi_v(R_o),\phi_t(\pi_{o})\right)$")
    add_body(doc, METHOD["grounding"][2])
    add_equation(doc, "6",
        r"$g^{\mathrm{con}}_{\mathrm{rel}}(b_s,b_o)=\frac{|\,b_s\cap b_o\,|}{\min(|b_s|,\,|b_o|)}$")
    add_body(doc, METHOD["grounding"][3])
    add_equation(doc, "7",
        r"$g^{\mathrm{dir}}_{\mathrm{rel}}(b_s,b_o)=\frac{\langle\, c_s-c_o,\; u_r\,\rangle}"
        r"{\|c_s-c_o\|}\;\in[-1,1]$")
    add_body(doc, METHOD["grounding"][4])

    add_heading(doc, "3.4. Detect-then-revise mitigation", 2)
    add_body(doc, METHOD["revise"][0])
    add_equation(doc, "8",
        r"$F(I,y)=\{\,o\in M_{\mathrm{obj}}(y)\;:\;s_{\mathrm{obj}}(I,o)<\tau_o\,\}$")
    add_body(doc, METHOD["revise"][1])
    add_equation(doc, "9",
        r"$y'=\oplus_{\,k\,:\;O(S_k)\cap F=\emptyset}\; S_k$")
    add_body(doc, METHOD["revise"][2])
    add_equation(doc, "10",
        r"$\mathrm{E}\!\left[L\right]=\sum_{o\,\in\,O_{\mathrm{true}}(y)}"
        r"\mathrm{P}\!\left[\,O(S(o))\cap F\neq\emptyset\,\right]$")
    add_body(doc, METHOD["revise"][3])

    add_heading(doc, "3.5. Calibrated decision fusion", 2)
    add_body(doc, METHOD["fusion"][0])
    add_equation(doc, "11",
        r"$\mathrm{logit}\,\mathrm{P}\!\left(H_0\mid p,g\right)="
        r"\mathrm{logit}\,\mathrm{P}(H_0)+\ell_p+\ell_g$")
    add_body(doc, METHOD["fusion"][1])
    add_equation(doc, "12",
        r"$z(I,q)=\hat{p}_{\mathrm{yes}}+\lambda\,\hat{g},\qquad "
        r"\hat{p}_{\mathrm{yes}}=\frac{p_{\mathrm{yes}}-\mu_p}{\sigma_p},\;\;"
        r"\hat{g}=\frac{g-\mu_g}{\sigma_g}$")
    add_body(doc, METHOD["fusion"][2])
    add_equation(doc, "13",
        r"$\tau^\star=Q_{1-r_0}(z),\qquad \hat{y}=\mathbf{1}\!\left[\,z\geq\tau^\star\,\right]$")
    add_body(doc, METHOD["fusion"][3])

    add_heading(doc, "3.6. Why redundant evidence cannot help: the residual view", 2)
    add_body(doc, METHOD["residual"][0])
    add_equation(doc, "14",
        r"$g=\alpha\,b+\varepsilon_g,\qquad \ell_p=\beta\,b+\varepsilon_p,\qquad "
        r"\mathrm{Cov}(g,\ell_p)=\alpha\beta\,\mathrm{Var}(b)$")
    add_body(doc, METHOD["residual"][1])
    add_equation(doc, "15",
        r"$g^{\perp}=g-\frac{\mathrm{Cov}(g,\ell_p)}{\mathrm{Var}(\ell_p)}\;\ell_p,"
        r"\qquad \mathrm{Var}(g^{\perp})=\left(1-\rho^2\right)\mathrm{Var}(g)$")
    add_body(doc, METHOD["residual"][2])
    add_body(doc, METHOD["residual"][3])

    add_heading(doc, "3.7. Deployment as a plug-in wrapper, and cost", 2)
    add_body(doc, METHOD["deploy"][0])
    add_body(doc, METHOD["deploy"][1])
    add_body(doc, METHOD["cost"])

    # ---- 4 Data ----
    add_heading(doc, "4. Data", 1)
    add_body(doc, DATA["intro"])
    add_heading(doc, "4.1. Source corpora", 2)
    add_body(doc, DATA["corpora"])
    add_heading(doc, "4.2. Evaluation benchmarks", 2)
    add_body(doc, DATA["bench"])
    add_table(doc,
              [("Table 2. ", "b"),
               ("Data used in this work. “Balance” is the fraction of positive "
                "(grounded) items, which matters because AMBER-relation is not balanced. "
                "HalluProbe-VL is released with this paper.", "")],
              ["Dataset", "What it scores", "Image source", "n", "Balance", "Used in"],
              TABLE2_ROWS, colwidths_mm=[28, 34, 26, 18, 16, 20],
              note="For HalluProbe-VL the released probe counts are 4,594 / 1,996 / 2,000; "
                   "detection AUC is computed on the evaluable subsets (Section 5.3).")
    add_heading(doc, "4.3. HalluProbe-VL: construction and leak-free design", 2)
    for para in DATA["hprobe"]:
        add_body(doc, para)

    # ---- 5 Experiments ----
    add_heading(doc, "5. Experiments", 1)
    add_heading(doc, "5.1. Setup", 2)
    for para in EXP["setup"]:
        add_body(doc, para)

    add_heading(doc, "5.2. Detection quality", 2)
    add_body(doc, EXP["detect"][0])
    add_table(doc,
              [("Table 3. ", "b"),
               ("Hierarchical detection quality, computed at full scale on the "
                "per-item scores. P and R are precision and recall at the best-F1 "
                "threshold τ*. The last column is a learned probabilistic alignment we "
                "trained first (Appendix B); it is at chance, which is why the final "
                "method fuses frozen off-the-shelf signals instead.", "")],
              ["Layer", "Benchmark", "n", "Pos.", "AUC", "AP", "Best-F1",
               "P@F1", "R@F1", "τ*", "Learned align. AUC"],
              TABLE3_ROWS,
              colwidths_mm=[24, 20, 12, 10, 12, 12, 13, 12, 12, 12, 18])
    add_figure(doc, "fig3_detection.png", 190,
               [("Fig. 3. ", "b"),
                ("Detection curves for the three hierarchical layers. (a) ROC with AUC "
                 "in the legend; the dashed diagonal is chance. (b) Precision–recall "
                 "with average precision in the legend; dotted lines mark each "
                 "benchmark’s class prior, the chance level for precision.", "")])
    for para in EXP["detect"][1:]:
        add_body(doc, para)

    add_heading(doc, "5.3. Generalization to the leak-free HalluProbe-VL", 2)
    add_body(doc, EXP["hprobe"][0])
    add_table(doc,
              [("Table 4. ", "b"),
               ("Detection on the leak-free HalluProbe-VL. n is the number of "
                "evaluable probes: attribute and relation probes require a groundable "
                "region, so fewer than the released 1,996 / 2,000 are scorable.", "")],
              ["Kind", "Released n", "Evaluated n", "AUC", "AP", "Best-F1"],
              TABLE4_ROWS, colwidths_mm=[24, 22, 24, 16, 16, 18])
    add_figure(doc, "fig4_halluprobe.png", 85,
               [("Fig. 4. ", "b"),
                ("Detection AUC per kind on HalluProbe-VL. Object detection on the "
                 "held-out COCO-2017 set matches its POPE score, showing the detector "
                 "generalizes rather than over-fitting one benchmark.", "")])
    add_body(doc, EXP["hprobe"][1])

    add_heading(doc, "5.4. The detection–mitigation gap", 2)
    for para in EXP["gap"]:
        add_body(doc, para)
    add_figure(doc, "fig5_gap.png", 190,
               [("Fig. 5. ", "b"),
                ("The detection–mitigation gap. (a) Mean presence score of true vs. "
                 "hallucinated caption objects; a region Grounding-DINO score (AUC 0.85) "
                 "separates them far better than a global CLIP score (AUC 0.73). "
                 "(b) Upper bound of detect-then-revise with the precise detector: "
                 "CHAIR-i falls while almost all true objects are kept. (c) Change in "
                 "CHAIR-i per strategy (lower is better); only detect-then-revise helps.", "")])

    add_heading(doc, "5.5. Object mitigation", 2)
    add_body(doc, EXP["mitig"][0])
    add_table(doc,
              [("Table 5. ", "b"),
               ("Object caption hallucination on CHAIR-500 (LLaVA-1.5-7B). CHAIR is "
                "lower-is-better; object recall is higher-is-better.", "")],
              ["Method", "CHAIR-i", "CHAIR-s", "Obj. recall"],
              TABLE5_ROWS, colwidths_mm=[46, 22, 22, 24])
    add_figure(doc, "fig6_mitigation.png", 140,
               [("Fig. 6. ", "b"),
                ("Object mitigation on CHAIR-500 (LLaVA-1.5). (a) Hallucination rates "
                 "for vanilla and the two revise variants. (b) Caption quality; BLEU-4 "
                 "and ROUGE-L are unchanged, so the drop in hallucination does not cost "
                 "quality.", "")])
    add_body(doc, EXP["mitig"][1])

    add_heading(doc, "5.6. Cross-backbone generalization of mitigation", 2)
    add_body(doc, EXP["xbb"][0])
    add_table(doc,
              [("Table 6. ", "b"),
               ("Cross-backbone object de-hallucination on CHAIR-500. CHAIR-i (%, "
                "lower is better); relative drop of the best variant in parentheses.", "")],
              ["Backbone", "Vanilla", "Self-rewrite", "Sent. removal", "(rel.)"],
              TABLE6_ROWS, colwidths_mm=[34, 20, 24, 24, 18])
    add_figure(doc, "fig7_crossbackbone.png", 140,
               [("Fig. 7. ", "b"),
                ("Cross-backbone object de-hallucination (CHAIR-i). Sentence removal "
                 "lowers hallucination on all four backbones; self-rewrite helps the "
                 "instruction-following ones but not InstructBLIP.", "")])
    add_body(doc, EXP["xbb"][1])

    add_heading(doc, "5.7. Discriminative decision fusion", 2)
    add_body(doc, EXP["discrim"][0])
    add_table(doc,
              [("Table 7. ", "b"),
               ("Cross-backbone POPE object discrimination with the yes-ratio matched "
                "to each model’s own, on a balanced 2,000-question subset. Higher is "
                "better.", "")],
              ["Backbone", "Vanilla Acc. (%)", "+grounding Acc. (%)", "Δ Acc.", "Δ F1"],
              TABLE7_ROWS, colwidths_mm=[34, 30, 34, 16, 16])
    add_figure(doc, "fig8_discrim_xbb.png", 140,
               [("Fig. 8. ", "b"),
                ("Cross-backbone discriminative gain on POPE at a matched yes-ratio. "
                 "Fusing grounding raises accuracy most on the weaker backbones and "
                 "leaves the strongest (LLaVA-1.6) unchanged in discrete accuracy, as "
                 "the residual view of Section 3.6 predicts.", "")])
    add_body(doc, EXP["discrim"][1])

    # ---- 6 Discussion ----
    add_heading(doc, "6. Discussion and limitations", 1)
    for para in DISCUSSION:
        add_body(doc, para)

    # ---- 7 Conclusion ----
    add_heading(doc, "7. Conclusion", 1)
    for para in CONCLUSION:
        add_body(doc, para)

    # ---- Appendix A ----
    add_heading(doc, "Appendix A. Single-backbone discriminative details", 1)
    add_body(doc, APPENDIX_A[0])
    add_table(doc,
              [("Table 8. ", "b"),
               ("POPE-9000 object discrimination (LLaVA-1.5-7B) with the yes-ratio "
                "matched to the model’s own. Higher is better.", "")],
              ["Method", "Acc. (%)", "F1 (%)", "Yes-ratio", "AUC"],
              TABLE8_ROWS, colwidths_mm=[36, 20, 20, 22, 18])
    add_figure(doc, "fig9_discriminative.png", 85,
               [("Fig. 9. ", "b"),
                ("POPE-9000 accuracy and F1 for vanilla vs. grounding fusion at a "
                 "matched yes-ratio of 0.460. The gain is a true improvement, not a "
                 "threshold shift.", "")])

    # ---- Appendix B ----
    add_heading(doc, "Appendix B. Negative results", 1)
    for para in APPENDIX_B:
        add_body(doc, para)

    # ---- References ----
    add_heading(doc, "References", 1)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(f"[{i}] {ref}")
        r.font.name = ARIAL; r.font.size = Pt(8.5)
        r._element.rPr.rFonts.set(qn("w:eastAsia"), ARIAL)

    out = os.path.join(HERE, "CMPSA_InformationFusion.docx")
    try:
        doc.save(out)
    except PermissionError:
        # the canonical file is open (locked) in Word -- save alongside it instead
        out = os.path.join(HERE, "CMPSA_InformationFusion_NEW.docx")
        doc.save(out)
        print("NOTE: canonical .docx is locked (open in Word?) -- saved to _NEW instead;")
        print("      close Word and rename, or re-run build_docx.py after closing.")
    print("saved ->", out)


# ================================================================= CONTENT
HIGHLIGHTS = [
    "A single training-free framework detects object, attribute, and relation hallucinations.",
    "Detect-then-revise cuts caption object hallucination by 20–41% on four backbones.",
    "A detection–mitigation gap is identified: shared co-occurrence bias breaks decoder guidance.",
    "A residual analysis explains when fused grounding evidence can and cannot help.",
    "A leak-free diagnostic, HalluProbe-VL, is released for three-kind detection.",
]

ABSTRACT = (
    "Multimodal large language models (MLLMs) often describe things that are not in the "
    "image. These errors fall into three kinds: made-up objects, wrong attributes, and "
    "wrong relations. We present a simple, training-free framework that fuses grounding "
    "evidence to handle all three kinds in one place. First, a hierarchical detector scores "
    "each kind with a matched grounding signal: zero-shot image–text similarity for objects, "
    "region-level contrast for attributes, and box geometry for relations. On public "
    "benchmarks it reaches an area under the ROC curve (AUC) of 0.80 for objects, 0.73 for "
    "attributes, and 0.71/0.62 for two relation settings, whereas a learned probabilistic "
    "alignment we trained first stays at chance (0.47–0.54). Second, we use the object "
    "detector to detect and then revise, dropping caption content that grounding reports as "
    "absent by a model self-rewrite or a backbone-agnostic sentence removal; this lowers the "
    "CHAIR-i rate on every backbone we test (by 20–41% with sentence removal) while keeping "
    "the true content, and clearly beats visual contrastive decoding. Third, we report a "
    "detection–mitigation gap: strong detection does not automatically give strong "
    "mitigation, because the grounding model shares the co-occurrence bias of the MLLM. We "
    "formalize this with a residual analysis showing that the useful part of an external "
    "source is what remains after projecting out the shared bias; this explains why "
    "token-by-token decoder guidance fails while a precise post-hoc revision succeeds. "
    "Fourth, a calibrated late fusion of grounding with the model’s own answer probability "
    "raises POPE accuracy by up to 2.0 points at a matched yes-ratio across four backbones, "
    "and — as the residual view predicts — helps least on the strongest backbone. The "
    "framework is training-free and plug-and-play: it deploys as a post-hoc wrapper around an "
    "unmodified backbone, its lightest level needing only the image and the output text — so "
    "it can wrap even a closed model served through an API — and it transfers across "
    "LLaVA-1.5, LLaVA-1.6, InstructBLIP, and Qwen-VL. We also release HalluProbe-VL, a "
    "leak-free diagnostic for three-kind detection, and we report the limits of the approach "
    "honestly."
)

KEYWORDS = ("multimodal large language models; hallucination; information fusion; "
            "visual grounding; training-free correction; image captioning")

INTRO = [
    "Multimodal large language models (MLLMs), such as LLaVA [1], describe images and answer "
    "visual questions with striking fluency. Yet they frequently state things that the image "
    "does not support [2]. Such hallucinations make the models hard to trust in careful "
    "settings such as medical reading or assisted perception, and they are now a central "
    "obstacle to deployment.",

    "Hallucinations come in three kinds. An object hallucination invents a thing that is not "
    "present (“a dog” in a dog-free scene). An attribute hallucination keeps the object but "
    "gives it the wrong property (“a red car” that is in fact blue). A relation hallucination "
    "gets the objects and attributes right but places them wrongly (“the cup on the table” "
    "when it is under it). Object errors are the most studied [3, 4]; attribute and relation "
    "errors are harder and less covered [5, 6].",

    "Most existing methods either detect hallucinations or mitigate them, and most focus on "
    "objects. Decoding-time methods such as visual contrastive decoding [7], over-trust "
    "penalties [8], and no-image contrast [9] reduce object errors during generation but are "
    "object-centric and sensitive to the generation setting. Detect-then-correct methods "
    "[10, 11, 12] repair the text after the fact, and training-based alignment [13] retrains "
    "the model on preference data. A natural hope is that a strong detector should yield a "
    "strong mitigator. We show that this hope holds only in part, and we explain why.",

    "We take an information-fusion view. Rather than trusting one signal, we fuse "
    "complementary grounding evidence at three levels of granularity, and we fuse that "
    "evidence with the model’s own answer where a decision is needed. This view yields a "
    "single, training-free framework (Fig. 1) that detects all three kinds with matched "
    "signals, revises captions where the object evidence is decisive, and makes yes/no "
    "decisions by a calibrated late fusion. Because every step consumes only the backbone’s "
    "inputs and outputs, the framework deploys as a plug-in wrapper around an unmodified — "
    "even closed — model (Section 3.7). Crucially, the fusion view also predicts when "
    "the approach must fail: an external source adds nothing where it is redundant with the "
    "model it is meant to correct. We make that precise and verify it.",

    "We make five contributions, each backed by full-scale, re-checked numbers. (1) A "
    "hierarchical detector that scores object, attribute, and relation hallucinations with a "
    "matched grounding signal, using only frozen off-the-shelf models. (2) A training-free, "
    "plug-and-play detect-then-revise mitigation for objects that lowers CHAIR-i on every "
    "backbone we test (by 20–41% with sentence removal) and beats visual contrastive "
    "decoding. (3) An analysis of the detection–mitigation gap, formalized by a residual "
    "decomposition that explains why decoder guidance fails while a precise post-hoc rewrite "
    "succeeds. (4) A calibrated decision-fusion rule that improves discriminative answering "
    "at a matched yes-ratio across four backbones. (5) HalluProbe-VL, a leak-free diagnostic "
    "for three-kind detection. We are also open about the limits: our mitigation is measured "
    "for objects, and the learned probabilistic alignment we first tried is not the reason "
    "our detectors work.",
]

RELATED = [
    "Benchmarks. CHAIR [4] measures object hallucination in captions against the objects in "
    "MS-COCO [16]. POPE [3] asks balanced yes/no questions about object presence. AMBER [5] "
    "adds attribute and relation questions and a caption task. HallusionBench [6] and similar "
    "suites test broader failures. We use these benchmarks and add a leak-free diagnostic "
    "built from COCO [16] and Visual Genome [17].",

    "Decoding-time mitigation. Visual contrastive decoding (VCD) [7] subtracts the logits of "
    "a noised image to reduce the language prior; OPERA [8] adds an over-trust penalty during "
    "beam search; M3ID [9] grounds by contrasting with a no-image branch; adaptive "
    "focal-contrast decoding [14] and contrastive-learning objectives [15] further curb "
    "object errors. These are effective but object-centric and, as we show, their effect can "
    "be sensitive to the generation setting.",

    "Detect-then-correct. Woodpecker [10] and LURE [11] detect object hallucinations and then "
    "correct the text, while Volcano [12] revises through self-feedback. Our object "
    "mitigation follows this paradigm; our differences are that we place it inside one "
    "hierarchical detection framework, that we compare grounding sources head to head, and "
    "that we give a clear account — with a supporting analysis — of when such correction "
    "works and when decoder guidance fails.",

    "Training-based alignment and fusion. A separate line retrains the model on human or AI "
    "preference data to reduce hallucination [13]. We instead keep every model frozen and "
    "fuse off-the-shelf grounding signals: CLIP [18] for image–text and region–text "
    "similarity and Grounding-DINO [19] for open-set detection. Table 1 positions our work "
    "against representative methods: prior training-free methods are object-centric and act "
    "at decoding time or correct only objects; none reports a unified detector for all three "
    "kinds, a head-to-head study of grounding sources, or the detection–mitigation gap that "
    "we make central, and none demonstrates the same mitigation across four different "
    "backbones.",
]

METHOD = {
    "intro":
        "Fig. 2 shows the structure of the algorithm. The design follows one rule: score each granularity with "
        "the evidence best suited to it, then fuse. This section formalizes the detector as a "
        "hypothesis test (Section 3.2), defines the three grounding signals (Section 3.3), "
        "derives the revision and decision rules (Sections 3.4–3.5), and then derives the "
        "condition under which fused evidence can help at all (Section 3.6) — the result that "
        "explains our central empirical finding.",
    "formulation": [
        "Given an image I and a prompt, an MLLM produces an output y that is either a caption "
        "or a yes/no answer. We read from y a set of claims M(y) that decomposes by granularity,",
        "where an object claim is a noun o, an attribute claim is a pair (o, a) with property "
        "a, and a relation claim is a triple (s, r, o). Each claim c carries a latent state: "
        "H₀, the claim is grounded in I, or H₁, the claim is hallucinated. A grounding "
        "function g(c) ∈ ℝ scores how well I supports c. All grounding models are frozen; no "
        "component is trained.",
    ],
    "hypothesis": [
        "Detection is a binary hypothesis test between H₀ (grounded) and H₁ (hallucinated). "
        "Given the score g(c), the Neyman–Pearson optimal rule thresholds the log-likelihood "
        "ratio,",
        "with η set by the tolerated false-alarm rate. We do not have the class-conditional "
        "densities p(g | H·), so we do not evaluate Eq. (2) directly. Instead we rely on a "
        "mild and testable assumption: the likelihood ratio is nondecreasing in g, i.e. the "
        "score family has the monotone likelihood ratio property. Under it, thresholding Λ is "
        "equivalent to thresholding g itself,",
        "which is exactly the detector we deploy, and it justifies the practice of reporting "
        "threshold-free summaries. Because any single τ is arbitrary, we report the area "
        "under the ROC curve (AUC), average precision (AP), and the best F1 over a full "
        "threshold sweep. Under the monotone assumption, AUC equals the probability that a "
        "grounded claim outranks a hallucinated one, so it measures the evidence quality of "
        "the signal itself rather than of one operating point. Section 5.2 reports the "
        "operating point (τ*, precision, recall) as well, because the gap analysis turns on "
        "precision rather than on AUC.",
    ],
    "grounding": [
        "Object level (OLD). For an object name o, the presence score is the cosine similarity "
        "between the image and a templated prompt π_o = “a photo of a o” in the shared CLIP "
        "space [18],",
        "where ϕ_v and ϕ_t are the CLIP image and text encoders. For correction we also use "
        "the maximum detection confidence s_obj(I, o) of Grounding-DINO [19] for the query o, "
        "which is region-level and, as Section 5.4 shows, decisively more precise. Attribute "
        "level (ALD). A global similarity cannot localize a property, so for a claim “o is a” "
        "we crop the region R_o given by the Grounding-DINO box for o and take a contrast "
        "between the attributed and the bare prompt,",
        "so the score measures the evidence for the attribute inside the object’s own region, "
        "with the object’s identity contribution differenced away. A low or negative score "
        "means the attribute is not supported where the object actually is. Relation level "
        "(RLD). Relations are geometric, so we score them from the two boxes b_s and b_o "
        "rather than from any similarity. Contact relations (“on”, “holding”, …) use the "
        "overlap ratio",
        "which is near 1 when the smaller box lies inside the larger and near 0 when the two "
        "are disjoint, and is deliberately asymmetric in scale so that a small object resting "
        "on a large one still scores high. Directional relations (“above”, “left of”, …) use "
        "the projection of the normalized centroid offset onto the relation’s unit axis u_r,",
        "where c_s and c_o are box centroids; the score is +1 when the pair is arranged "
        "exactly as stated and −1 when it is exactly reversed. A claim is flagged when the "
        "geometry disagrees with the stated relation. Using a distinct, physically meaningful "
        "signal per level is the core of the fusion: each granularity is scored by the "
        "evidence suited to it, rather than by one global similarity that is blind to "
        "location and geometry.",
    ],
    "revise": [
        "The detector drives a training-free repair. For a caption y we flag the objects whose "
        "precise, region-level score falls below a per-object threshold,",
        "If F is empty we keep y unchanged. Otherwise we revise y into y′ with one of two "
        "operators. Self-rewrite asks the same MLLM to rewrite y with a short instruction "
        "listing F as “not in the image”; the output stays fluent but the operator needs a "
        "backbone that follows instructions. Sentence removal keeps only the sentences whose "
        "object set does not intersect F,",
        "where S_k are the sentences of y, O(S_k) is the object set of sentence S_k, and ⊕ "
        "concatenates. Sentence removal is deterministic and backbone-agnostic. Its cost is a "
        "collateral recall loss: a true object is dropped whenever it happens to share a "
        "sentence with a flagged one, so the expected loss L is",

        "where S(o) is the sentence containing o. Eq. (10) makes the trade-off explicit and "
        "bounds it: the loss is driven by co-occurrence within a sentence, not by the number "
        "of flags, so it stays small when hallucinated objects are mentioned in dedicated "
        "clauses and grows when the caption bundles many objects per sentence. It also shows "
        "why the operator is safe at a precise threshold: a smaller F shrinks every term. We "
        "measure this cost directly (object recall 1.83→1.73 in Section 5.5) rather than "
        "assuming it away. Both operators are training-free and touch only the captions that "
        "carry a flag — about one image in six in our runs.",
    ],
    "fusion": [
        "For yes/no questions we fuse the model’s own evidence with grounding. Treat the "
        "answer probability p_yes and the object grounding g as two sources bearing on the "
        "same hypothesis. If the sources were conditionally independent given the hypothesis, "
        "the posterior log-odds would decompose additively,",
        "with ℓ_p and ℓ_g the log-likelihood-ratio contributions of the two sources. Eq. (11) "
        "is the classical justification for linear late fusion: adding calibrated evidence in "
        "log-odds space is Bayes-optimal under conditional independence. We therefore "
        "standardize each source on a held-out calibration split and add them with a single "
        "weight,",
        "where μ and σ are the calibration mean and standard deviation of each source; "
        "z-scoring is the affine calibration that makes the two scales commensurable, and λ "
        "— the only free parameter — absorbs the unknown ratio of the two sources’ "
        "informativeness and is fitted on the calibration split. To avoid the trivial "
        "“answer no more often” shortcut, we then set the decision threshold so that the "
        "fused yes-ratio equals the model’s own yes-ratio r₀,",
        "where Q is the quantile function of z. This constrains the fused rule to the same "
        "base rate as the raw model. Consequently the decision is a fixed-quantile cut of z, "
        "and its accuracy is a monotone function of how well z ranks positives above "
        "negatives — that is, of the fused AUC. Any accuracy gain at the matched yes-ratio is "
        "therefore evidence of added ranking information rather than of a shifted operating "
        "point. We confirm empirically that the fused AUC rises whenever the accuracy rises "
        "(Section 5.7).",
    ],
    "residual": [
        "Eq. (11) assumes the two sources are conditionally independent. They are not. The "
        "grounding model and the MLLM are trained on overlapping web image–text data and "
        "inherit the same object co-occurrence statistics: both are inclined to believe a "
        "keyboard accompanies a monitor. Write b for that shared co-occurrence bias and model "
        "each source as a loading on b plus an idiosyncratic term,",
        "with ε_g and ε_p uncorrelated with b and with each other, and write ρ = ρ(g, ℓ_p) for "
        "the induced correlation. The component of the grounding score that the model does "
        "not already contain is the residual after projecting g onto ℓ_p,",
        "Only g⊥ can change a decision that ℓ_p alone would get wrong; the collinear part αb "
        "merely re-states what the model already believes. Hence the achievable gain from "
        "fusing this source scales with Var(g⊥) = (1 − ρ²)·Var(g), and vanishes as ρ → 1 no "
        "matter how large Var(g) is. A source can be simultaneously accurate and useless.",
        "This single quantity organizes the whole paper. (i) A global CLIP score is strongly "
        "collinear with the MLLM’s language prior — precisely on the high-co-occurrence "
        "objects that get hallucinated — so its ρ is high, its residual is small, and steering "
        "the decoder with it cannot remove those hallucinations (Section 5.4). (ii) A "
        "region-level Grounding-DINO score answers a different question (“is there evidence "
        "at this location?”), so its residual is large and a post-hoc revision driven by it "
        "works. (iii) On a backbone whose own ℓ_p is already strong, Var(g⊥) buys little "
        "headroom, which predicts the null result we observe on LLaVA-1.6 (Section 5.7). The "
        "practical lesson for fusion is that source precision on the model’s error region, "
        "not marginal source accuracy, is what must be maximized.",
    ],
    "deploy": [
        "Nothing in Sections 3.2–3.5 reads the backbone’s weights, gradients, hidden "
        "features, or decoding state. Detection consumes the image and the output text; "
        "revision consumes the text and the flag set; decision fusion consumes one scalar per "
        "question. The framework therefore deploys as a plug-in wrapper — a layer that sits "
        "between an unmodified backbone and its consumer, takes the pair (image, output) in, "
        "and returns a verified caption or a calibrated answer — rather than as a change to "
        "the model. What the host system must expose determines which capabilities the "
        "wrapper offers, and we make this explicit as three access levels.",

        "L0 — image and output text only. Hierarchical detection and sentence removal operate "
        "here; they can wrap any backbone, including a closed model served through an API "
        "that returns nothing but text. L1 — the ability to prompt the backbone once more. "
        "Self-rewrite operates here, and additionally needs a backbone that follows "
        "instructions. L2 — first-token probabilities. The calibrated decision fusion needs "
        "p_yes, which open-weight deployments expose; we do not claim L2 for API-only models "
        "that hide token probabilities. These levels are a description of what was measured, "
        "not an aspiration: the L0 sentence removal transfers to all four backbones "
        "(Table 6), the L1 self-rewrite fails on exactly the one backbone that follows "
        "instructions poorly, and every L2 fusion result was run with logit access "
        "(Table 7).",
    ],
    "cost":
        "The framework adds no training and modest inference cost. Detection costs one CLIP "
        "text encoding per claim and one image encoding per image, both cached; attribute and "
        "relation claims add one Grounding-DINO pass per image to obtain boxes, plus one CLIP "
        "pass per crop. Mitigation touches only flagged captions — about one image in six in "
        "our runs — so self-rewrite adds a second MLLM generation on that fraction only, and "
        "sentence removal is a string operation with no model call. Decision fusion adds one "
        "CLIP text encoding per question and a scalar update; the calibration of λ and τ* is a "
        "one-off quantile computation on a held-out split.",
}

DATA = {
    "intro":
        "This section describes every corpus the paper uses: the sources the images and "
        "annotations come from, the evaluation benchmarks and their sizes and class balance, "
        "and the construction of HalluProbe-VL, the diagnostic we release. Table 2 summarizes "
        "them.",
    "corpora":
        "Images and scene graphs come from two public corpora. MS-COCO [16] supplies natural "
        "photographs with exhaustive instance annotations over 80 everyday categories; we use "
        "the 2014 validation split, on which POPE and CHAIR are defined, and the 2017 "
        "validation split (5,000 images), which we reserve for our leak-free diagnostic. "
        "Visual Genome [17] supplies dense region descriptions, attributes, and subject–"
        "relation–object triples, and is the only large source with the relation annotations "
        "our RLD layer needs. We use these corpora strictly as sources of ground truth: no "
        "image is used for training, since no component of the method is trained.",
    "bench":
        "We evaluate on four public benchmarks and one we release. POPE [3] poses balanced "
        "yes/no questions about object presence over COCO-2014 images (9,000 questions across "
        "its random, popular, and adversarial subsets), and is our object-detection and "
        "discrimination bed. AMBER [5] contributes attribute questions (4,764) and contact-"
        "relation questions (1,664). Because no public benchmark covers directional relations "
        "at scale, we build a Visual Genome probe set, VG-Rel (2,000 items), by sampling "
        "annotated directional triples and their reversed counterfactuals. CHAIR [4] measures "
        "object hallucination in free-form captions against COCO ground-truth objects; "
        "following common practice we score 500 COCO-2014 validation images and report "
        "CHAIR-i (fraction of mentioned object instances that are hallucinated) and CHAIR-s "
        "(fraction of captions containing at least one hallucination), together with object "
        "recall so that a method cannot win by saying less. Class balance is 50% positives "
        "except for AMBER-relation, which is 58.6% positive; we report it in Table 2 because "
        "it makes that row’s F1 look better than its AUC warrants.",
    "hprobe": [
        "Existing detection beds have two problems for our purpose. They are object-centric, "
        "and their images overlap the data the community has tuned on. We therefore release "
        "HalluProbe-VL, a diagnostic with balanced probes for all three kinds, built to be "
        "leak-free and to keep the evaluated tools out of the labelling loop.",

        "Object probes are drawn from COCO-2017 val, which is disjoint from the COCO-2014 val "
        "images used by POPE and CHAIR; this is what “leak-free” means here — a detector "
        "tuned on the 2014 beds cannot have seen these images in that role. Each image "
        "contributes a positive probe for an annotated category and a negative probe for an "
        "absent category, giving a 50/50 balance over 4,594 probes. Negatives are not sampled "
        "uniformly: we draw them from the categories that most frequently co-occur with the "
        "image’s true categories in the COCO training statistics, which makes them hard "
        "negatives aimed exactly at the co-occurrence bias of Section 3.6. Attribute (1,996) "
        "and relation (2,000) probes are built the same way from Visual Genome, pairing each "
        "annotated attribute or triple with a counterfactual that alters the property or "
        "reverses the relation.",

        "Two design choices matter for validity. First, every label comes from the source "
        "annotations (COCO instances, VG attributes and relationships), never from a tool "
        "prediction, so CLIP and Grounding-DINO remain the evaluated detectors rather than "
        "the labellers — otherwise the benchmark would measure agreement with itself. Second, "
        "we release the probes with relative image references rather than redistributing "
        "images, so the set inherits the source licences. A caveat we state plainly: "
        "attribute and relation probes are only scorable when the referenced object can be "
        "localized to a region, so the AUC in Section 5.3 is computed on the evaluable "
        "subsets (555 and 634 probes) rather than on all released probes; the object layer "
        "has no such restriction and is scored on all 4,594.",
    ],
}

EXP = {
    "setup": [
        "Backbones and tools. Unless stated otherwise the MLLM is LLaVA-1.5-7B [1] with the "
        "CLIP-ViT-L/14-336 encoder [18]. Grounding uses the same CLIP and Grounding-DINO [19]. "
        "For the cross-backbone study we add LLaVA-1.6-7B [22], InstructBLIP-7B [20], and "
        "Qwen-VL-Chat [21]. All models run in half precision on a single 24 GB GPU and every "
        "model is frozen.",
        "Protocol. Captions are decoded greedily with a fixed 80-token budget; the "
        "detect-then-revise threshold is τ_o = 0.30 throughout. Caption quality uses "
        "BLEU-4 [23], ROUGE-L [24], and object coverage. Discrimination is scored at a "
        "matched yes-ratio with λ and τ* fitted on a held-out half of the questions and "
        "applied to the other half. All numbers are computed at full scale and re-checked; "
        "small pilot runs proved unstable and are not used for headline claims.",
    ],
    "detect": [
        "Table 3 and Fig. 3 give the detector results. All three kinds are well above chance: "
        "AUC 0.804 (object), 0.729 (attribute), and 0.706/0.622 (relation, contact/direction), "
        "with average precision tracking AUC. The ordering AUC(object) > AUC(attribute) > "
        "AUC(relation) mirrors both the difficulty of the three kinds and the precision of "
        "the matched signal: object presence is the cleanest signal, while directional "
        "relations rest on fragile geometry and are the weakest layer.",

        "The comparison in the last column of Table 3 is the reason the method is "
        "training-free. Before settling on frozen signals we trained a probabilistic "
        "cross-modal alignment intended to be the grounding engine; on the same benchmarks it "
        "scores 0.507, 0.543, 0.473, and 0.500 AUC — chance, at every level, including below "
        "chance for contact relations. We diagnose its representational collapse in "
        "Appendix B. We report it because it is the honest counterfactual to the design: the "
        "gains here come from fusing well-matched off-the-shelf evidence, not from a learned "
        "alignment.",

        "The operating point matters more than the summary numbers, and Table 3 makes it "
        "explicit. At the best-F1 threshold every layer sits in a high-recall, "
        "low-precision regime — object detection reaches recall 0.836 at precision 0.662, and "
        "directional relations recall 0.913 at precision 0.541. In other words, tuned for F1, "
        "these detectors flag nearly every hallucination but also flag many grounded claims. "
        "That is acceptable for reporting and filtering, where a human or a downstream stage "
        "adjudicates, but it is exactly the wrong operating point for editing text "
        "automatically: at precision 0.66 a third of the removals would delete true content. "
        "This is the first sign of the gap studied next, and it is why the revision stage "
        "does not reuse the best-F1 threshold but moves to a precise, high-threshold "
        "operating point of a different (region-level) source.",

        "The precision–recall panel, Fig. 3(b), also shows why AMBER-relation’s best F1 "
        "(0.795) looks strong next to its modest AUC (0.706): its class prior is 0.586, so "
        "the chance level for precision is already high. Comparing each curve against its own "
        "dotted prior — rather than across benchmarks — is the fair reading.",
    ],
    "hprobe": [
        "A detector that only works on the bed it was tuned on is of little use, so we "
        "re-evaluate on the leak-free HalluProbe-VL of Section 4.3. Table 4 and Fig. 4 report "
        "the result.",
        "The object layer reaches AUC 0.828 on held-out COCO-2017 images with "
        "high-co-occurrence hard negatives, essentially matching its POPE score (0.804). "
        "Since the probes are disjoint from POPE’s images and are adversarially chosen "
        "against co-occurrence, this is evidence that the object detector generalizes rather "
        "than over-fitting one benchmark. Attribute (0.685) and relation (0.651) transfer at "
        "a lower level, consistent with their weaker public-benchmark scores and with the "
        "smaller evaluable subsets; we do not claim more for them than Table 4 shows.",
    ],
    "gap": [
        "A strong detector need not give a strong mitigator. Section 3.6 predicts exactly "
        "when it will not: when the source is collinear with the model it corrects. We test "
        "that prediction using COCO captions from LLaVA-1.5-7B and summarize it in Fig. 5.",
        "Grounding shares the model’s bias. Fig. 5(a) shows the mean presence score of "
        "caption objects, split into true and hallucinated. With a global CLIP score the two "
        "groups overlap heavily (0.55 vs. 0.49, AUC 0.73): the objects the MLLM invents are "
        "high-co-occurrence objects that CLIP also rates as likely present. This is a direct "
        "measurement of a large ρ in Eq. (14) — a small residual. A region-level "
        "Grounding-DINO score separates the same two groups far better (0.63 vs. 0.38, AUC "
        "0.85), because it asks a question the language prior cannot answer: is there "
        "evidence at this location?",
        "Guiding the decoder token by token fails. We suppressed the tokens of absent objects "
        "during generation, naively and with per-class calibration. Both keep recall but fail "
        "to lower CHAIR (Fig. 5(c), first two bars); the calibrated variant even leaves "
        "CHAIR-i slightly worse. Visual contrastive decoding, under a matched setting, does "
        "not help either. All three inject a source whose residual is small precisely where "
        "the errors are.",
        "Precise, post-hoc grounding succeeds. Using the region-level detector in a post-hoc "
        "rewrite gives a clean trade-off: raising the threshold removes more hallucinations "
        "while keeping almost all true objects (Fig. 5(b)); at threshold 0.30 the upper bound "
        "removes 35% of hallucinations while keeping 97.8% of true objects. With a global "
        "CLIP score, reaching the same CHAIR drop costs about 20% of true objects; with the "
        "precise detector it costs about 2%. The deciding factor is not the mechanism’s name "
        "but the source’s precision on the model’s error region — Var(g⊥), not accuracy.",
    ],
    "mitig": [
        "Table 5 and Fig. 6 give the CHAIR-500 results on LLaVA-1.5-7B. Self-rewrite lowers "
        "CHAIR-i from 16.9% to 14.1% (−16.9% relative) and sentence removal lowers it further "
        "to 11.7% (−30.9%); CHAIR-s falls in step. Self-rewrite keeps object recall almost "
        "unchanged (1.83→1.82 true objects per caption) and does not hurt caption quality "
        "(BLEU-4 and ROUGE-L are unchanged; Fig. 6(b)); sentence removal is stronger on "
        "hallucination but a little coarser on recall (1.83→1.73), the collateral loss "
        "predicted by Eq. (10).",
        "Under a matched setting, visual contrastive decoding did not help — it slightly "
        "raised CHAIR-i, lost more recall, and its effect changes with the generation length "
        "— so our revision is both stronger and more stable.",
    ],
    "xbb": [
        "The method is backbone-agnostic: detection scores any model’s output, and sentence "
        "removal edits its text without training. Table 6 and Fig. 7 apply the same pipeline "
        "to four backbones of different design. Sentence removal lowers CHAIR-i on all four "
        "by 20–41% relative, keeping most true objects.",
        "Self-rewrite helps the three instruction-following backbones (up to −16.9%) but not "
        "InstructBLIP, an older model that follows the rewrite instruction poorly. This is a "
        "clear sign that the self-rewrite variant needs a capable backbone, whereas the "
        "deterministic sentence-removal variant does not and is therefore the safer default. "
        "Notably the largest relative gain (−40.7%) is on InstructBLIP, the backbone that "
        "hallucinates most — the pattern the residual view predicts, since a weaker model "
        "leaves a larger error region for an external source to correct. In deployment terms "
        "(Section 3.7) sentence removal is the L0 capability: it asks nothing of the backbone "
        "beyond its text, which is exactly why it transfers everywhere.",
    ],
    "discrim": [
        "Beyond captions, the fused evidence improves yes/no answering. Table 7 and Fig. 8 "
        "fuse the same CLIP grounding with each backbone’s own answer probability on a "
        "balanced 2,000-question POPE subset at a matched yes-ratio (Eqs. 12–13). Accuracy "
        "rises on three of the four backbones, and most on the weaker ones (InstructBLIP "
        "+2.0, Qwen-VL +1.4, LLaVA-1.5 +1.3 points).",
        "On LLaVA-1.6, whose own answers are already the strongest (answer AUC 0.951), the "
        "fusion adds no discrete gain at the matched yes-ratio, though its ranking AUC still "
        "improves slightly (0.951→0.956). This null is not a failure of the fusion rule but "
        "the behaviour Section 3.6 predicts: when ℓ_p is already strong, the residual "
        "Var(g⊥) has little left to correct. Across the four backbones the gain tracks the "
        "backbone’s own weakness, which is what a genuine external source should do. The "
        "single-backbone POPE-9000 breakdown, where accuracy rises from 86.69% to 87.49% at a "
        "matched yes-ratio, is reported in Appendix A.",
    ],
}

DISCUSSION = [
    "Why fusion is the right frame. Every step is a fusion of complementary evidence: three "
    "granularity-specific signals fuse for detection, a precise source fuses with the caption "
    "for revision, and the model’s answer fuses with grounding for a decision. The "
    "detection–mitigation gap shows that fusion must respect the dependence between "
    "sources. The residual analysis of Section 3.6 turns this into an actionable design rule: "
    "do not ask whether a source is accurate, ask how much of it the model does not already "
    "know. That single quantity accounts for three otherwise unrelated observations — the "
    "failure of decoder guidance, the success of region-level post-hoc revision, and the null "
    "result on the strongest backbone.",

    "Practicality. The wrapper view of Section 3.7 makes the framework a drop-in safety "
    "layer for systems that already exist: a captioning service can route every (image, "
    "caption) pair through detection and sentence removal without touching its serving stack "
    "(L0), and a deployment with logit access can add the calibrated fusion (L2). Only "
    "flagged captions — about one in six — incur a second model call, no component needs a "
    "GPU beyond what the backbone already uses, and attaching a new backbone amounts to "
    "implementing a small caption/answer interface, which is how the four backbones in "
    "Section 5.6 were added. The released code ships this as a pip-installable wrapper "
    "(hgfusion) whose L0 path is a three-line call — construct the wrapper, pass any "
    "model’s caption, receive the verified text — with the L1 rewrite exposed as a "
    "user-supplied callback and the L2 fusion as a calibrate-then-answer pair.",

    "Limitations. We state them plainly. (i) Our measured mitigation is for objects; CHAIR "
    "and AMBER’s caption task both score objects only. When we extended the rewrite to "
    "attributes and relations we did not obtain a measured gain: parsing attribute and "
    "relation claims from free text is noisy, box overlap misses attached-but-non-touching "
    "parts, and there is no standard caption metric for these kinds. We therefore keep the "
    "attribute and relation contribution at the detection level, where it is measured "
    "(Tables 3–4), and treat hierarchical mitigation as future work. (ii) The "
    "detect-then-revise paradigm follows prior work [10, 11, 12]; our value is the single "
    "hierarchical framework, the head-to-head grounding study, and the gap analysis with its "
    "residual account. (iii) A learned probabilistic alignment we first tried is at chance "
    "(Table 3) and is not part of the final method. (iv) The self-rewrite variant needs a "
    "capable backbone; on InstructBLIP only sentence removal works, and sentence removal can "
    "drop a true object that shares a sentence with a hallucinated one. (v) Relation "
    "detection is our weakest layer (AUC 0.622 on directional probes); box geometry cannot "
    "resolve relations that depend on depth or on function rather than position.",
]

CONCLUSION = [
    "We presented a training-free framework that fuses grounding evidence to detect object, "
    "attribute, and relation hallucinations in MLLMs, to mitigate object hallucination by a "
    "detect-then-revise step guided by a precise detector, and to answer yes/no questions by "
    "a calibrated decision fusion. The framework reaches useful detection AUC on all three "
    "kinds — where a learned alignment stays at chance — generalizes to a leak-free "
    "diagnostic, and across four different backbones cuts caption object hallucination by "
    "20–41% without hurting quality while raising discriminative accuracy by up to 2.0 points "
    "at a matched yes-ratio. Because it consumes only the backbone’s inputs and outputs, it "
    "attaches to an existing system as a post-hoc wrapper whose lightest level needs nothing "
    "but the image and the output text. Our central message is the detection–mitigation gap and its "
    "explanation: what a fused source contributes is not its accuracy but its residual — the "
    "part the model does not already believe. Because grounding and the MLLM share a "
    "co-occurrence bias, decoder guidance fails while a precise post-hoc fusion works, and "
    "the same quantity predicts where the method helps least. We hope the released "
    "HalluProbe-VL and this honest account help future work close the gap for attributes and "
    "relations.",
]

APPENDIX_A = [
    "Table 8 and Fig. 9 report the single-backbone discriminative result behind Section 5.7. "
    "The yes-ratio is held at the model’s own value (0.460), so the accuracy and F1 gains "
    "cannot come from answering “no” more often. The fused ranking AUC also rises "
    "(0.936→0.941), so the gain is added ranking information, consistent with the argument "
    "in Section 3.5.",
]

APPENDIX_B = [
    "The learned alignment. Before adopting frozen signals we trained a probabilistic "
    "cross-modal alignment with optimal-transport matching, intended to be the grounding "
    "engine of the framework. It collapsed: with free means and a uniform marginal the "
    "matching objective drove all embeddings toward a single point, so retrieval was near "
    "chance and the resulting detectors scored 0.47–0.54 AUC (Table 3). We repaired the "
    "objective — correspondence loss, matched-pair Gaussian penalty, debiased divergence, and "
    "a variance hinge — which fixed the collapse and restored retrieval far above chance, but "
    "the repaired alignment still did not beat the off-the-shelf signals at detection. We "
    "report it because it is the counterfactual that motivates the training-free design.",

    "Token-level suppression. Suppressing absent-object tokens during generation, naively and "
    "with per-class calibration, did not lower CHAIR (it slightly raised CHAIR-i). Section 3.6 "
    "explains why: the suppressed source is collinear with the language prior on exactly the "
    "objects that get hallucinated.",

    "Hierarchical revision. A rewrite that also edited attributes and relations did not "
    "improve object CHAIR (+0.002 versus −22.8% for object-only revision in the same "
    "setting). Three causes compounded: a mixed correction instruction led the model to "
    "cosmetic edits rather than decisive removal; relation grounding produced false positives "
    "on attached but non-overlapping parts (a stop sign “on” its pole); and no quantitative "
    "caption benchmark exists for attribute or relation hallucination, so a gain could not be "
    "measured even where it occurred. Both negative results point to source precision and "
    "reliable claim parsing as the binding constraints.",
]

# ---------------------------------------------------------------- tables
TABLE1_ROWS = [
    ["VCD [7]",         "Decoding contrast",   "✓", "–", "–", "–", "✓", "✓", "–", "–"],
    ["OPERA [8]",       "Decoding penalty",    "✓", "–", "–", "–", "✓", "✓", "–", "–"],
    ["M3ID [9]",        "Decoding contrast",   "✓", "–", "–", "–", "✓", "✓", "–", "–"],
    ["Woodpecker [10]", "Detect-then-correct", "✓", "~", "–", "✓", "✓", "✓", "–", "–"],
    ["LURE [11]",       "Detect-then-correct", "✓", "–", "–", "✓", "✓", "–", "–", "–"],
    ["Volcano [12]",    "Self-feedback revise","✓", "~", "–", "✓", "✓", "–", "–", "–"],
    ["Ours",            "Detect-then-revise + fusion", "✓", "~", "~", "✓", "✓", "✓", "✓", "✓"],
]
TABLE2_ROWS = [
    ["POPE [3]",        "object presence (yes/no)",      "COCO 2014 val", "9,000",  "50%",   "§5.2, §5.7"],
    ["AMBER-attr [5]",  "attribute claims",              "AMBER",         "4,764",  "50%",   "§5.2"],
    ["AMBER-rel [5]",   "contact relations",             "AMBER",         "1,664",  "58.6%", "§5.2"],
    ["VG-Rel (ours)",   "directional relations",         "Visual Genome", "2,000",  "50%",   "§5.2"],
    ["CHAIR-500 [4]",   "caption object hallucination",  "COCO 2014 val", "500 img","–",     "§5.4–5.6"],
    ["HalluProbe-VL",   "object / attribute / relation", "COCO 2017 val + VG",
     "4,594 / 1,996 / 2,000", "50%", "§5.3"],
]
TABLE3_ROWS = [
    ["OLD (object)",          "POPE",       "9,000", "50%",   "0.804", "0.821", "0.739", "0.662", "0.836", "0.136",  "0.507"],
    ["ALD (attribute)",       "AMBER-attr", "4,764", "50%",   "0.729", "0.706", "0.717", "0.601", "0.888", "−0.013", "0.543"],
    ["RLD (rel., contact)",   "AMBER-rel",  "1,664", "58.6%", "0.706", "0.726", "0.796", "0.683", "0.952", "0.066",  "0.473"],
    ["RLD (rel., direction)", "VG-Rel",     "2,000", "50%",   "0.622", "0.577", "0.679", "0.541", "0.913", "−0.127", "0.500"],
]
TABLE4_ROWS = [
    ["Object",    "4,594", "4,594", "0.828", "0.854", "0.749"],
    ["Attribute", "1,996", "555",   "0.685", "0.660", "0.689"],
    ["Relation",  "2,000", "634",   "0.651", "0.591", "0.695"],
]
TABLE5_ROWS = [
    ["Vanilla",                 "16.9", "28.6", "1.83"],
    ["Self-rewrite (ours)",     "14.1", "24.2", "1.82"],
    ["Sentence removal (ours)", "11.7", "19.0", "1.73"],
]
TABLE6_ROWS = [
    ["LLaVA-1.5-7B",    "16.9", "14.1", "11.7", "−30.9%"],
    ["LLaVA-1.6-7B",    "7.5",  "7.1",  "6.0",  "−19.9%"],
    ["InstructBLIP-7B", "19.2", "19.3", "11.4", "−40.7%"],
    ["Qwen-VL-Chat",    "15.9", "13.8", "10.6", "−33.3%"],
]
TABLE7_ROWS = [
    ["LLaVA-1.5-7B",    "87.0", "88.3", "+1.3", "+1.4"],
    ["LLaVA-1.6-7B",    "87.9", "87.9", "+0.0", "+0.0"],
    ["InstructBLIP-7B", "84.6", "86.6", "+2.0", "+2.1"],
    ["Qwen-VL-Chat",    "85.1", "86.5", "+1.4", "+1.6"],
]
TABLE8_ROWS = [
    ["Vanilla",    "86.69", "86.12", "0.460", "0.936"],
    ["+grounding", "87.49", "86.96", "0.460", "0.941"],
]

# ---------------------------------------------------------------- references
REFERENCES = [
    "H. Liu, C. Li, Q. Wu, and Y. J. Lee, “Visual instruction tuning,” in Advances in Neural Information Processing Systems (NeurIPS), 2023.",
    "Z. Bai, P. Wang, T. Xiao, T. He, Z. Han, Z. Zhang, and M. Z. Shou, “Hallucination of multimodal large language models: A survey,” arXiv:2404.18930, 2024.",
    "Y. Li, Y. Du, K. Zhou, J. Wang, W. X. Zhao, and J.-R. Wen, “Evaluating object hallucination in large vision-language models,” in Proc. Conf. Empirical Methods in Natural Language Processing (EMNLP), 2023.",
    "A. Rohrbach, L. A. Hendricks, K. Burns, T. Darrell, and K. Saenko, “Object hallucination in image captioning,” in Proc. Conf. Empirical Methods in Natural Language Processing (EMNLP), 2018.",
    "J. Wang, Y. Wang, G. Xu, J. Zhang, Y. Gu, H. Jia, M. Yan, J. Zhang, and J. Sang, “An LLM-free multi-dimensional benchmark for MLLMs hallucination evaluation (AMBER),” arXiv:2311.07397, 2023.",
    "T. Guan, F. Liu, X. Wu, R. Xian, Z. Li, X. Liu, X. Wang, L. Chen, F. Huang, Y. Yacoob, D. Manocha, and T. Zhou, “HallusionBench: An advanced diagnostic suite for entangled language hallucination and visual illusion in large vision-language models,” in Proc. IEEE/CVF Conf. Computer Vision and Pattern Recognition (CVPR), 2024.",
    "S. Leng, H. Zhang, G. Chen, X. Li, S. Lu, C. Miao, and L. Bing, “Mitigating object hallucinations in large vision-language models through visual contrastive decoding,” in Proc. IEEE/CVF Conf. Computer Vision and Pattern Recognition (CVPR), 2024.",
    "Q. Huang, X. Dong, P. Zhang, B. Wang, C. He, J. Wang, D. Lin, W. Zhang, and N. Yu, “OPERA: Alleviating hallucination in multi-modal large language models via over-trust penalty and retrospection-allocation,” in Proc. IEEE/CVF Conf. Computer Vision and Pattern Recognition (CVPR), 2024.",
    "F. Favero, L. Zancato, M. Trager, S. Choudhary, P. Perera, A. Achille, A. Swaminathan, and S. Soatto, “Multi-modal hallucination control by visual information grounding (M3ID),” in Proc. IEEE/CVF Conf. Computer Vision and Pattern Recognition (CVPR), 2024.",
    "S. Yin, C. Fu, S. Zhao, T. Xu, H. Wang, D. Sui, Y. Shen, K. Li, X. Sun, and E. Chen, “Woodpecker: Hallucination correction for multimodal large language models,” Science China Information Sciences, 2024.",
    "Y. Zhou, C. Cui, J. Yoon, L. Zhang, Z. Deng, C. Finn, M. Bansal, and H. Yao, “Analyzing and mitigating object hallucination in large vision-language models (LURE),” in Proc. Int. Conf. Learning Representations (ICLR), 2024.",
    "S. Lee, S. H. Park, Y. Jo, and M. Seo, “Volcano: Mitigating multimodal hallucination through self-feedback guided revision,” in Proc. Conf. North American Chapter of the Association for Computational Linguistics (NAACL), 2024.",
    "T. Yu, Y. Yao, H. Zhang, T. He, Y. Han, G. Cui, J. Hu, Z. Liu, H.-T. Zheng, M. Sun, and T.-S. Chua, “RLHF-V: Towards trustworthy MLLMs via behavior alignment from fine-grained correctional human feedback,” in Proc. IEEE/CVF Conf. Computer Vision and Pattern Recognition (CVPR), 2024.",
    "Z. Chen, Z. Zhao, H. Luo, H. Yao, B. Li, and J. Zhou, “HALC: Object hallucination reduction via adaptive focal-contrast decoding,” in Proc. Int. Conf. Machine Learning (ICML), 2024.",
    "C. Jiang, H. Xu, M. Dong, J. Chen, W. Ye, M. Yan, Q. Ye, J. Zhang, F. Huang, and S. Zhang, “Hallucination augmented contrastive learning for multimodal large language models,” in Proc. IEEE/CVF Conf. Computer Vision and Pattern Recognition (CVPR), 2024.",
    "T.-Y. Lin, M. Maire, S. Belongie, J. Hays, P. Perona, D. Ramanan, P. Dollár, and C. L. Zitnick, “Microsoft COCO: Common objects in context,” in Proc. European Conf. Computer Vision (ECCV), 2014.",
    "R. Krishna, Y. Zhu, O. Groth, J. Johnson, K. Hata, J. Kravitz, S. Chen, Y. Kalantidis, L.-J. Li, D. A. Shamma, M. S. Bernstein, and L. Fei-Fei, “Visual Genome: Connecting language and vision using crowdsourced dense image annotations,” International Journal of Computer Vision, vol. 123, no. 1, pp. 32–73, 2017.",
    "A. Radford, J. W. Kim, C. Hallacy, A. Ramesh, G. Goh, S. Agarwal, G. Sastry, A. Askell, P. Mishkin, J. Clark, G. Krueger, and I. Sutskever, “Learning transferable visual models from natural language supervision (CLIP),” in Proc. Int. Conf. Machine Learning (ICML), 2021.",
    "S. Liu, Z. Zeng, T. Ren, F. Li, H. Zhang, J. Yang, C. Li, J. Yang, H. Su, J. Zhu, and L. Zhang, “Grounding DINO: Marrying DINO with grounded pre-training for open-set object detection,” in Proc. European Conf. Computer Vision (ECCV), 2024.",
    "W. Dai, J. Li, D. Li, A. M. H. Tiong, J. Zhao, W. Wang, B. Li, P. Fung, and S. Hoi, “InstructBLIP: Towards general-purpose vision-language models with instruction tuning,” in Advances in Neural Information Processing Systems (NeurIPS), 2023.",
    "J. Bai, S. Bai, S. Yang, S. Wang, S. Tan, P. Wang, J. Lin, C. Zhou, and J. Zhou, “Qwen-VL: A versatile vision-language model for understanding, localization, text reading, and beyond,” arXiv:2308.12966, 2023.",
    "H. Liu, C. Li, Y. Li, B. Li, Y. Zhang, S. Shen, and Y. J. Lee, “LLaVA-NeXT: Improved reasoning, OCR, and world knowledge,” technical report, 2024.",
    "K. Papineni, S. Roukos, T. Ward, and W.-J. Zhu, “BLEU: A method for automatic evaluation of machine translation,” in Proc. Annual Meeting of the Association for Computational Linguistics (ACL), 2002.",
    "C.-Y. Lin, “ROUGE: A package for automatic evaluation of summaries,” in Text Summarization Branches Out, ACL Workshop, 2004.",
    "Z. Sun, S. Shen, S. Cao, H. Liu, C. Li, Y. Shen, C. Gan, L.-Y. Gui, Y.-X. Wang, Y. Yang, K. Keutzer, and T. Darrell, “Aligning large multimodal models with factually augmented RLHF,” arXiv:2309.14525, 2023.",
    "P. Wang, S. Bai, S. Tan, S. Wang, Z. Fan, J. Bai, K. Chen, X. Liu, J. Wang, W. Ge, Y. Fan, K. Dang, M. Du, X. Ren, R. Men, D. Liu, C. Zhou, J. Zhou, and J. Lin, “Qwen2-VL: Enhancing vision-language model’s perception of the world at any resolution,” arXiv:2409.12191, 2024.",
]


if __name__ == "__main__":
    build()
